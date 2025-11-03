import os
import datetime
from flask import Flask, render_template, request, send_file, flash, redirect
from docx import Document

# Inisialisasi Flask
app = Flask(__name__, template_folder="ui")
app.secret_key = 'secretkey'

TEMPLATE_PATH = os.path.join("template", "template.docx")
OUTPUT_DIR = "outputs"
os.makedirs(OUTPUT_DIR, exist_ok=True)

# Placeholder standar
EXPECTED_FIELDS = [
    "nomor_surat",
    "nomor_polis",
    "nama_tertanggung",
    "unit",
    "tkp",
    "alamat_tertanggung",
    "alamat_sopir",
    "alamat_bengkel",
    "kantor_polisi",
    "dol"
]

# Mapping untuk dropdown asuransi
ASURANSI_INFO = {
    "TOB": """Kepada Yth. :
Bapak Rizqi Abdul Ghani
PT ASURANSI TOTAL BERSAMA 
Citra Tower, 27th Floor
Jl. Benyamin Suaeb Blok A6 RT 13/ RW 06
Kebon Kosong Kec. Kemayoran – Jakarta Pusat 10630""",
    "SOMPO": """Kepada Yth. :
Bapak Arief Hariyanto
Head Of Motor Claim Departement
PT. Sompo Insurance Indonesia
Mayapada Tower II, 19th floor
Jl. Jend. Sudirman Kav. 27 Jakarta 12920""",
    "ACA": """Kepada Yth. :
Bapak Agus Suryono
PT ASURANSI CENTRAL ASIA
Gedung Hermina Tower 1 (Lantai 3)
Jl. HBR Motik Blok B-10 Kemayoran, Jakarta 10610"""
}

SURVEY_FEE_INFO = {
    "TOB": "",        
    "SOMPO": """Survey fee									Rp    400.000
Terbilang ( empat ratus ribu rupiah )""", 
    "ACA": ""
}

PERUSAHAAN_INFO = {
    "TOB": "PT ASURANSI TOTAL BERSAMA",
    "SOMPO": "PT. SOMPO INSURANCE INDONESIA",
    "ACA": "PT ASURANSI CENTRAL ASIA"
}

DIREKTUR_INFO = {
    "TOB": "Rizqi Abdul Ghani",
    "SOMPO": """Arief Hariyanto
Administration Head							Head of Motor Claim Departement""",
    "ACA": "Agus Suryono"
}

# ---------- Fungsi Replace Placeholder ---------- #

def replace_in_paragraph(paragraph, mapping):
    full_text = ''.join(run.text for run in paragraph.runs)
    replaced = full_text
    changed = False
    for key, val in mapping.items():
        if key in replaced:
            replaced = replaced.replace(key, val)
            changed = True
    if changed:
        for i in range(len(paragraph.runs)-1, -1, -1):
            paragraph._p.remove(paragraph.runs[i]._r)
        paragraph.add_run(replaced)

def replace_in_table(table, mapping):
    for row in table.rows:
        for cell in row.cells:
            replace_in_block(cell, mapping)

def replace_in_block(block, mapping):
    for p in block.paragraphs:
        replace_in_paragraph(p, mapping)
    for tbl in getattr(block, "tables", []):
        replace_in_table(tbl, mapping)

def replace_placeholders(doc: Document, mapping: dict):
    replace_in_block(doc, mapping)
    for table in doc.tables:
        replace_in_table(table, mapping)
    for section in doc.sections:
        try:
            replace_in_block(section.header, mapping)
        except Exception:
            pass
        try:
            replace_in_block(section.footer, mapping)
        except Exception:
            pass

# ---------- Route Utama ---------- #

@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        # Ambil data form
        data = {key: request.form.get(key, "").strip() for key in EXPECTED_FIELDS}

        # Mapping placeholder dari form
        mapping = {f"{{{k}}}": v for k, v in data.items()}

        # Tambahkan mapping untuk placeholder "kepada_asuransi"
        asuransi_key = request.form.get("asuransi")
        mapping["{kepada_asuransi}"] = ASURANSI_INFO.get(asuransi_key, "")
        
        survey_key = request.form.get("asuransi")
        mapping["{survey_fee}"] = SURVEY_FEE_INFO.get(survey_key, "")

        perusahaan_key = request.form.get("asuransi")
        mapping["{Perusahaan}"] = PERUSAHAAN_INFO.get(perusahaan_key, "")

        direktur_key = request.form.get("asuransi")
        mapping["{direktur}"] = DIREKTUR_INFO.get(direktur_key, "")

        # Pastikan template ada
        if not os.path.exists(TEMPLATE_PATH):
            flash(f"Template tidak ditemukan: {TEMPLATE_PATH}")
            return redirect(request.url)

        try:
            doc = Document(TEMPLATE_PATH)
            replace_placeholders(doc, mapping)
        except Exception as e:
            flash(f"Terjadi kesalahan saat mengganti placeholder: {e}")
            return redirect(request.url)

        # Simpan hasil
        safe_nomor = data.get("nomor_surat", "") or "hasil"
        timestamp = datetime.datetime.now().strftime("%Y%m%d%H%M%S")
        out_name = f"{safe_nomor}_{timestamp}.docx".replace(" ", "_")
        out_path = os.path.join(OUTPUT_DIR, out_name)
        doc.save(out_path)

        flash("✅ Dokumen berhasil dibuat dan sedang diunduh.")
        return send_file(out_path, as_attachment=True, download_name=out_name)

    return render_template("index.html")

if __name__ == "__main__":
    app.run(debug=True)
