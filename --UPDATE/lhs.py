import os
import datetime
from flask import Flask, render_template, request, send_file, flash, redirect
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from num2words import num2words

# Inisialisasi Flask
app = Flask(__name__, template_folder="./UI")
app.secret_key = 'secretkey'

TEMPLATE_PATH = os.path.join("Master", "LHS_MASTER.docx")
OUTPUT_DIR = "Downloads/LHS_Output"
os.makedirs(OUTPUT_DIR, exist_ok=True)

# Placeholder standar
EXPECTED_FIELDS = [
    "asuransi",
    "tertanggung",
    "polis",
    "unit",
    "plat",
    "case",
    "sta",
    "datesta",
    "stksc",
    "datestksc",
    "surveyor",
    "alamat_tertanggung",
    "lapos",
    "rangka",
    "mesin",
    "dol",
    "tkp",
    "doltime",
    "nama_tertanggung",
    "alamat_tertanggung",
    "usia_tertanggung",
    "pekerjaan",
    "keterangan_tertanggung",
    "data"
]

# Mapping untuk dropdown asuransi
ASURANSI_INFO = {
    "TOB": "PT. Asuransi Total Bersama",
    "SOMPO": "PT. Sompo Insurance Indonesia",
    "ACA": "PT. Asuransi Central Asia"
}

CASE_INFO = {
    "kecelakaan": "kecelakaan",
    "kehilangan": "kehilangan",
}

SURVEYOR_INFO = {
    "albert": "Albertus S.",
    "yulius": "Yulius B.",
    "eko": "Eko",
    "fredy": "Fredy",
    "samino": "Samino",
}

# ---------- Fungsi Replace Placeholder ---------- #

def replace_in_paragraph(paragraph, mapping):
    full_text = ''.join(run.text for run in paragraph.runs)
    replaced = full_text
    changed = False
    for key, val in mapping.items():
        if key in replaced:
            replaced = replaced.replace(key, val)
            changed = True
    if changed:
        for i in range(len(paragraph.runs)-1, -1, -1):
            paragraph._p.remove(paragraph.runs[i]._r)
        paragraph.add_run(replaced)

def replace_in_table(table, mapping):
    for row in table.rows:
        for cell in row.cells:
            replace_in_block(cell, mapping)

def replace_in_block(block, mapping):
    for p in block.paragraphs:
        replace_in_paragraph(p, mapping)
    for tbl in getattr(block, "tables", []):
        replace_in_table(tbl, mapping)

def replace_placeholders(doc: Document, mapping: dict):
    replace_in_block(doc, mapping)
    for table in doc.tables:
        replace_in_table(table, mapping)
    for section in doc.sections:
        try:
            replace_in_block(section.header, mapping)
        except Exception:
            pass
        try:
            replace_in_block(section.footer, mapping)
        except Exception:
            pass

# ---------- Route Utama ---------- #

@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        # Ambil data form
        data = {key: request.form.get(key, "").strip() for key in EXPECTED_FIELDS}

        # Mapping placeholder dari form
        mapping = {f"{{{k}}}": v for k, v in data.items()}

        # Tambahkan mapping untuk placeholder "asuransi"
        asuransi_key = request.form.get("asuransi")
        mapping["{asuransi}"] = ASURANSI_INFO.get(asuransi_key, "")

        case_key = request.form.get("case")
        mapping["{case}"] = CASE_INFO.get(case_key, "")

        surveyor_key = request.form.get("surveyor")
        mapping["{surveyor}"] = SURVEYOR_INFO.get(surveyor_key, "")

        titles = request.form.getlist("receipt_title[]")
        amounts = request.form.getlist("receipt_amount[]")
        receipts = []
        for t, a in zip(titles, amounts):
            if t.strip() or a.strip():
                cleaned = "".join(ch for ch in a if ch.isdigit())
                amt = int(cleaned) if cleaned else 0
                receipts.append({"title": t.strip(), "amount": amt})

        

        # Pastikan template ada
        if not os.path.exists(TEMPLATE_PATH):
            flash(f"Template tidak ditemukan: {TEMPLATE_PATH}")
            return redirect(request.url)

        try:
            doc = Document(TEMPLATE_PATH)
            replace_placeholders(doc, mapping)
        except Exception as e:
            flash(f"Terjadi kesalahan saat mengganti placeholder: {e}")
            return redirect(request.url)

        # Simpan hasil
        safe_nomor = data.get("asuransi", "") or "hasil"
        timestamp = datetime.datetime.now().strftime("%Y%m%d%H%M%S")
        out_name = f"{safe_nomor}_{timestamp}.docx".replace(" ", "_")
        out_path = os.path.join(OUTPUT_DIR, out_name)
        try:
            doc.save(out_path)
        except Exception as e:
            flash(f"Gagal menyimpan file: {e}")
            return redirect(request.url)

        flash("✅ Dokumen berhasil dibuat dan sedang diunduh.")
        return send_file(out_path, as_attachment=True, download_name=out_name)

    return render_template("lhs.html")

if __name__ == "__main__":
    app.run(debug=True)
